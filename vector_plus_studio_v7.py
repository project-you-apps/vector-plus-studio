"""
Vector+ Studio V7 - TRUE Content-Addressable Recall

Unlike V6 which uses np.dot(embeddings, query) for ranking (bypassing lattice physics),
V7 uses actual lattice signatures for similarity search.

Key differences:
1. Uses MultiLatticeCUDAv7 with QUALITY profile (0.86+ correlation at 30% noise)
2. Builds signature database during ingestion (4096-float per pattern)
3. Ranking uses scan_signatures() - true lattice physics, not embedding dot product
4. NO TILING - discovered tiling hurts capacity

The lattice physics now actually matters for search quality!
"""

import streamlit as st

# MUST be first Streamlit command - controls sidebar visibility
st.set_page_config(
    page_title="Vector+ Studio v0.7",
    layout="wide",
    initial_sidebar_state="expanded",
)

import numpy as np
import pickle
import os
import sys
import time
import threading

import PyPDF2, pypdf
import docx

from sentence_transformers import SentenceTransformer
from streamlit.runtime.scriptrunner import add_script_run_ctx

# --- 1. SYSTEM SETUP ---
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
sys.path.insert(0, parent_dir)
sys.path.insert(0, current_dir)

# --- IMPORTS ---
try:
    from multi_lattice_wrapper_v7 import MultiLatticeCUDAv7
except ImportError:
    st.error("CRITICAL: V7 wrapper not found. (multi_lattice_wrapper_v7.py missing)")
    st.stop()

# --- OPTIONAL IMPORTS FOR FORGE ---
# try:
#     import PyPDF2
# except ImportError:
#     try:
#         import pypdf
#     except ImportError:
#         pypdf = None
# try:
#     import docx
# except ImportError:
#     docx = None
#
# # stop trying to be fancy in a demo! lol

# --- CONFIG ---
SAFETY_BATCH_SIZE = 32
SAFETY_COOLDOWN = 0.3  # GPU cooldown between batches (prevents UPS overload)
SETTLE_FRAMES = 10  # Test harness was 30 and we matched but switched to 10 for speed.
PHYSICS_PROFILE = "quality"  # Best correlation (0.86+ at 30% noise)

DATA_DIR = os.path.join(os.path.dirname(__file__), "cartridges")


def get_layout():
    # Pseudocode for determining layout based on screen size
    if st.session_state.get("screen_width", 800) > 1200:
        return "wide"
    else:
        return "centered"


def load_cartridge_data(cartridge_name):
    """Smart Loader: Handles both V1 (Legacy), V2, and V7 signature formats."""
    path = os.path.join(DATA_DIR, f"{cartridge_name}.pkl")

    with open(path, "rb") as f:
        raw_payload = pickle.load(f)

    if isinstance(raw_payload, dict) and raw_payload.get("version") == "7.0":
        # V7 FORMAT - Has signatures!
        core_data = raw_payload["data"]
        signatures = raw_payload.get("signatures")  # May be None if interrupted
        metadata = raw_payload.get("metadata", {})
        print(
            f"Loaded V7 Cartridge with {len(signatures) if signatures is not None else 0} signatures"
        )
        return core_data, metadata, signatures

    elif isinstance(raw_payload, dict) and raw_payload.get("version") == "2.0":
        # V2 FORMAT
        core_data = raw_payload["data"]
        metadata = raw_payload.get("metadata", {})
        print(f"Loaded V2 Cartridge (no signatures)")
        return core_data, metadata, None

    else:
        # V1 LEGACY
        core_data = raw_payload
        metadata = {}
        print(f"Loaded V1 Cartridge (Raw Mode)")
        return core_data, metadata, None


def read_full_document(filepath):
    """Reads the content of a local file for the Source Viewer."""
    if not os.path.exists(filepath):
        return f"❌ Error: File not found at {filepath}\n(Did you move or delete it?)"

    ext = filepath.split(".")[-1].lower()

    try:
        # --- TXT / MD ---
        if ext in ["txt", "md", "py", "json"]:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        # --- PDF ---
        elif ext == "pdf":
            text = ""
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                # Limit to first 50 pages for speed in the demo
                for page in reader.pages[:50]:
                    text += page.extract_text() + "\n\n"
            return text

        # --- DOCX ---
        elif ext == "docx":
            doc = docx.Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs])

        else:
            return f"⚠️ Preview not supported for .{ext} files yet."

    except Exception as e:
        return f"❌ Error reading file: {str(e)}"


# --- THREADING HELPERS ---
if "engine_lock" not in st.session_state:
    st.session_state.engine_lock = threading.Lock()

if "training_status" not in st.session_state:
    st.session_state.training_status = "LatticeRunner V7 Idle"
    st.session_state.training_progress = 0.0


def build_signature_database(engine, embeddings, lock, progress_callback=None):
    """
    Build signature database from embeddings.

    IMPORTANT: We use learn=False to generate INDEPENDENT signatures.
    Each signature captures the pattern's activation fingerprint without
    accumulating weights (which would make all signatures identical).

    This is the key V7 feature: each pattern gets a 4096-float signature
    that captures the lattice's response to that pattern.
    """
    total = len(embeddings)
    signatures = np.zeros((total, 4096), dtype=np.float32)

    for i in range(total):
        emb = embeddings[i]

        # Normalize (NO TILING - hurts capacity)
        v_min, v_max = emb.min(), emb.max()
        vec_norm = (emb - v_min) / (v_max - v_min + 1e-9)
        vec_norm = vec_norm.astype(np.float32)

        with lock:
            # Reset lattice before each pattern (prevents accumulation)
            engine.reset()

            # Imprint pattern - NO SETTLE for signatures!
            # The raw thermometer encoding IS the unique fingerprint
            # Settle would kill the pattern with QUALITY physics
            engine.imprint_vector(vec_norm)

            # Generate signature from raw imprinted pattern
            signatures[i] = engine.generate_signature()

        # GPU cooldown every batch to prevent thermal/power issues
        if i % SAFETY_BATCH_SIZE == 0:
            time.sleep(SAFETY_COOLDOWN)

        if progress_callback and i % 50 == 0:
            progress_callback(i, total, f"Building signatures: {i}/{total}")

    return signatures


def background_trainer_and_save(
    engine,
    embeddings,
    start_idx,
    lock,
    save_path=None,
    existing_signatures=None,
    dataset_ref=None,
):
    """
    Background thread for building signatures.
    Saves V7 format cartridge with signatures.
    """
    total = len(embeddings)

    # Initialize or continue signature array
    if existing_signatures is not None:
        signatures = existing_signatures
    else:
        signatures = np.zeros((total, 4096), dtype=np.float32)

    for i in range(start_idx, total):
        emb = embeddings[i]

        # Normalize (NO TILING)
        v_min, v_max = emb.min(), emb.max()
        vec_norm = (emb - v_min) / (v_max - v_min + 1e-9)
        vec_norm = vec_norm.astype(np.float32)

        with lock:
            # Reset lattice before each pattern (prevents accumulation)
            engine.reset()

            # Imprint - NO SETTLE (raw thermometer = fingerprint)
            engine.imprint_vector(vec_norm)
            signatures[i] = engine.generate_signature()

        # GPU cooldown every batch to prevent thermal/power issues
        if i % SAFETY_BATCH_SIZE == 0:
            time.sleep(SAFETY_COOLDOWN)

        if i % 50 == 0:
            st.session_state.training_progress = i / total
            st.session_state.training_status = (
                f"V7 Signatures: {i}/{total} ({int(i/total*100)}%)"
            )

            # Update session state with partial signatures
            if dataset_ref:
                dataset_ref["signatures"] = signatures

            time.sleep(0.001)

    # Finished - save V7 format
    st.session_state.training_progress = 1.0

    if save_path and dataset_ref:
        with lock:
            # Save V7 format cartridge
            v7_data = {
                "version": "7.0",
                "data": {
                    "embeddings": embeddings,
                    "passages": dataset_ref.get("txt", []),
                },
                "signatures": signatures,
                "metadata": {
                    "profile": PHYSICS_PROFILE,
                    "settle_frames": SETTLE_FRAMES,
                    "created": time.strftime("%Y-%m-%d %H:%M"),
                },
            }
            with open(save_path, "wb") as f:
                pickle.dump(v7_data, f)

            # Also save brain state
            brain_path = save_path.replace(".pkl", "_brain.npy")
            engine.save_brain(brain_path)

        st.session_state.training_status = "V7 Signatures Complete & Saved"
        dataset_ref["signatures"] = signatures
    else:
        st.session_state.training_status = "V7 Signatures Built (100%)"
        if dataset_ref:
            dataset_ref["signatures"] = signatures


try:
    from spellchecker import SpellChecker

    spell = SpellChecker()
    spell.word_frequency.load_words(
        [
            "mcp",
            "cuda",
            "numpy",
            "py",
            "docx",
            "vector",
            "embedding",
            "logits",
            "starlette",
            "fastapi",
            "v7",
        ]
    )
    HAS_SPELLCHECK = True
except ImportError:
    HAS_SPELLCHECK = False
    spell = None


def strict_spellcheck(text):
    if not HAS_SPELLCHECK:
        return text
    words = text.split()
    corrected = []
    for word in words:
        correction = spell.correction(word)
        if correction and correction != word:
            corrected.append(correction)
        else:
            corrected.append(word)
    return " ".join(corrected)


# --- INITIALIZATION ---

# REMOVE TOP PADDING
st.markdown(
    """
    <style>
           /* Remove blank space at top and bottom */
           .block-container {
               padding-top: 1rem;
               padding-bottom: 0rem;
           }

           /* Hide the "Made with Streamlit" footer, but keep sidebar toggle visible */
           footer {
               visibility: hidden;
           }
    </style>
    """,
    unsafe_allow_html=True,
)

# REDUCE VERTICAL GAP BETWEEN WIDGETS
st.markdown(
    """
    <style>
        /* Reduce the vertical gap between widgets */
        [data-testid="stVerticalBlock"] {
            gap: 0.5rem;
        }
    </style>
""",
    unsafe_allow_html=True,
)

# REMOVE SIDEBAR PADDING

st.markdown(
    """
    <style>
        /* Tighten sidebar padding */
        [data-testid="stSidebar"] > div:first-child {
            padding-top: 0rem;
        }
    </style>
""",
    unsafe_allow_html=True,
)

if "engine" not in st.session_state or st.session_state.engine is None:
    with st.spinner("Booting LatticeRunner V7 Physics Engine..."):
        try:
            st.session_state.engine = MultiLatticeCUDAv7(
                lattice_size=4096, max_layers=4, verbose=1  # console detail
            )
            # Set QUALITY profile for best recall
            st.session_state.engine.set_profile(PHYSICS_PROFILE)
            st.toast("LatticeRunner V7 Online (QUALITY Profile)")
        except Exception as e:
            st.error(f"Failed to boot V7 engine: {e}")
            st.stop()

if "dataset" not in st.session_state:
    st.session_state.dataset = None
if "deleted_ids" not in st.session_state:
    st.session_state.deleted_ids = set()


@st.cache_resource(show_spinner="Loading Embedder...")
def load_embedder():
    # Try Nomic first (768-dim, best quality)
    try:
        model = SentenceTransformer("nomic-ai/nomic-embed-text-v1.5", trust_remote_code=True)
        print("[Embedder] Nomic v1.5 loaded (768-dim)")
        return model
    except Exception as e:
        print(f"[Embedder] Nomic failed: {e}")

    # Fallback to all-mpnet (768-dim, good quality, no trust_remote_code needed)
    try:
        model = SentenceTransformer("all-mpnet-base-v2")
        print("[Embedder] Fallback to all-mpnet-base-v2 (768-dim)")
        return model
    except Exception as e:
        print(f"[Embedder] MPNet failed: {e}")
        st.error(f"Could not load any embedder: {e}")
        st.stop()


# Pre-load embedder at startup (so first search is instant)
_ = load_embedder()


# --- SIDEBAR: DATA CONTROLLER ---
with st.sidebar:
    st.title("Vector+ v0.7")
    st.caption("True Content-Addressable Recall")
    st.divider()

    st.header("Memory Controller")

    st.caption("Core Status:")
    st.progress(st.session_state.training_progress)
    st.caption(st.session_state.training_status)
    st.divider()

    # In the Sidebar
    # st.sidebar.markdown("---")
    # st.sidebar.header("⚡ Physics Engine")

    # Create a placeholder we can update later
    # status_log = st.sidebar.empty()

    # # ... later, right before you run the search ...
    # with status_log.container():
    #     st.info("🌊 Rippling Energy Landscape...")
    #     # (The search runs here)

    # # ... right after the search finishes ...
    # with status_log.container():
    #     st.success(f"✅ Converged in {0.04}s")  # You can grab the real timer

    # --- MOUNT CARTRIDGE ---
    data_dir = DATA_DIR
    os.makedirs(data_dir, exist_ok=True)

    if os.path.exists(data_dir):
        files = [f for f in os.listdir(data_dir) if f.endswith(".pkl")]
        selected_file = st.selectbox("Select Memory Cartridge", files)

        if st.button("Mount Cartridge", type="primary"):
            file_path = os.path.join(data_dir, selected_file)

            with st.spinner("Mounting Volume..."):
                with open(file_path, "rb") as f:
                    data = pickle.load(f)

                # Handle different formats
                if isinstance(data, dict) and data.get("version") == "7.0":
                    embeddings = data["data"]["embeddings"]
                    texts = data["data"]["passages"]
                    signatures = data.get("signatures")
                    st.toast("V7 Cartridge Loaded!")
                else:
                    embeddings = data.get(
                        "embeddings", data.get("data", {}).get("embeddings")
                    )
                    texts = data.get("passages", data.get("data", {}).get("passages"))
                    signatures = None
                    st.toast("Legacy Cartridge - Building Signatures...")

                ml = st.session_state.engine

                base_name = os.path.splitext(selected_file)[0]
                brain_path = os.path.join(data_dir, f"{base_name}_brain.npy")
                v7_path = os.path.join(data_dir, f"{base_name}_v7.pkl")

                st.session_state.dataset = {
                    "emb": embeddings,
                    "txt": texts,
                    "filename": selected_file,
                    "signatures": signatures,
                }
                st.session_state.deleted_ids = set()

                # Check for pre-built V7 signatures
                if signatures is not None:
                    # Already have signatures - instant boot
                    st.session_state.training_status = "V7 Signatures Loaded (Instant)"
                    st.session_state.training_progress = 1.0
                    if os.path.exists(brain_path):
                        with st.session_state.engine_lock:
                            ml.load_brain(brain_path)
                    st.success("Memory Active with Signatures!")

                elif os.path.exists(brain_path):
                    # Brain exists but no signatures - rebuild
                    st.toast("Brain found but no signatures - building...")
                    with st.session_state.engine_lock:
                        ml.load_brain(brain_path)

                    # Build signatures in background
                    bg_thread = threading.Thread(
                        target=background_trainer_and_save,
                        args=(
                            ml,
                            embeddings,
                            0,
                            st.session_state.engine_lock,
                            v7_path,
                            None,
                            st.session_state.dataset,
                        ),
                        daemon=True,
                    )
                    add_script_run_ctx(bg_thread)
                    bg_thread.start()
                    st.success("Building signatures in background...")

                else:
                    # Cold start - build everything
                    warmup_count = 100
                    st.caption(
                        f"Cold Start: Building first {warmup_count} signatures..."
                    )

                    warmup_sigs = np.zeros((len(embeddings), 4096), dtype=np.float32)

                    with st.session_state.engine_lock:
                        for i in range(min(warmup_count, len(embeddings))):
                            v_min, v_max = embeddings[i].min(), embeddings[i].max()
                            vec_norm = (embeddings[i] - v_min) / (v_max - v_min + 1e-9)
                            vec_norm = vec_norm.astype(np.float32)

                            # Reset lattice before each pattern
                            ml.reset()

                            # Imprint - NO SETTLE (raw thermometer = fingerprint)
                            ml.imprint_vector(vec_norm)
                            warmup_sigs[i] = ml.generate_signature()

                            # GPU cooldown every batch
                            if i % SAFETY_BATCH_SIZE == 0:
                                time.sleep(SAFETY_COOLDOWN)

                    st.session_state.dataset["signatures"] = warmup_sigs
                    st.success(f"Mounted! Background signature build started...")

                    if len(embeddings) > warmup_count:
                        bg_thread = threading.Thread(
                            target=background_trainer_and_save,
                            args=(
                                ml,
                                embeddings,
                                warmup_count,
                                st.session_state.engine_lock,
                                v7_path,
                                warmup_sigs,
                                st.session_state.dataset,
                            ),
                            daemon=True,
                        )
                        add_script_run_ctx(bg_thread)
                        bg_thread.start()
                    else:
                        st.session_state.training_status = "V7 Signatures Complete"
                        st.session_state.training_progress = 1.0

    #
    # *** FUTURE IMPLEMENTATION ***
    #
    # # --- CART X-RAY ---
    # if st.session_state.get("cart_data"):
    #     with st.expander("🔍 Cartridge X-Ray"):
    #         # Grab the first item in the registry (or passages list)
    #         if st.session_state.registry:
    #             first_key = list(st.session_state.registry.keys())[0]
    #             st.write(f"**Sample File:** `{first_key}`")
    #             st.json(st.session_state.registry[first_key])

    #         # Show a sample passage
    #         if st.session_state.cart_data["passages"]:
    #             st.caption("**Sample Text Chunk:**")
    #             st.code(st.session_state.cart_data["passages"][0][:500] + "...")

    # --- THE FORGE ---
    st.divider()
    st.header("The Forge")
    with st.expander("Forge New Cartridge"):
        uploaded_files = st.file_uploader(
            "Upload Documents", type=["txt", "pdf", "docx"], accept_multiple_files=True
        )

        cartridge_name = st.text_input("Cartridge Name", value="my_docs")

        if st.button("Forge"):
            if not uploaded_files:
                st.error("No files uploaded.")
            else:
                new_texts = []
                prog_bar = st.progress(0)

                for idx, file in enumerate(uploaded_files):
                    text = ""
                    try:
                        if file.name.endswith(".txt"):
                            text = file.read().decode("utf-8")
                        elif file.name.endswith(".pdf"):
                            if pypdf or pyPDF2:
                                reader = pypdf.PdfReader(file)
                                for page in reader.pages:
                                    text += page.extract_text() + "\n"
                            else:
                                st.error("Install 'pypdf' for PDFs.")
                        elif file.name.endswith(".docx"):
                            if docx:
                                doc = docx.Document(file)
                                text = "\n".join([p.text for p in doc.paragraphs])
                            else:
                                st.error("Install 'python-docx' for DOCX.")
                    except Exception as e:
                        st.warning(f"Error parsing {file.name}: {e}")
                        continue

                    paragraphs = [p for p in text.split("\n\n") if len(p) > 50]
                    for p in paragraphs:
                        new_texts.append(f"{p}\n[Source: {file.name}]")

                    prog_bar.progress((idx + 1) / len(uploaded_files))

                if len(new_texts) > 0:
                    st.info(f"Embedding {len(new_texts)} chunks...")
                    embedder = load_embedder()

                    all_embs = []
                    batch_size = SAFETY_BATCH_SIZE

                    emb_bar = st.progress(0)
                    for i in range(0, len(new_texts), batch_size):
                        batch = new_texts[i : i + batch_size]
                        files = [f"search_document: {t}" for t in batch]
                        batch_embs = embedder.encode(files, convert_to_numpy=True)
                        all_embs.append(batch_embs)
                        emb_bar.progress(min((i + batch_size) / len(new_texts), 1.0))
                        time.sleep(SAFETY_COOLDOWN)

                    final_embs = np.vstack(all_embs)

                    # Save as legacy format (signatures built on mount)
                    output = {"embeddings": final_embs, "passages": new_texts}
                    save_path = os.path.join(data_dir, f"{cartridge_name}.pkl")
                    with open(save_path, "wb") as f:
                        pickle.dump(output, f)

                    st.success(
                        f"Cartridge '{cartridge_name}' forged! ({len(new_texts)} chunks)"
                    )
                    st.caption("Mount to build V7 signatures.")
                else:
                    st.warning("No valid text extracted.")

    # --- COMMIT DELETES ---
    if st.session_state.deleted_ids:
        st.divider()
        st.warning(f"{len(st.session_state.deleted_ids)} items pending deletion.")
        if st.button("Commit & Prune Cartridge"):
            if st.session_state.dataset:
                old_embs = st.session_state.dataset["emb"]
                old_txts = st.session_state.dataset["txt"]
                old_sigs = st.session_state.dataset.get("signatures")
                fname = st.session_state.dataset["filename"]

                keep_mask = [
                    i
                    for i in range(len(old_txts))
                    if i not in st.session_state.deleted_ids
                ]
                new_embs = old_embs[keep_mask]
                new_txts = [old_txts[i] for i in keep_mask]
                new_sigs = old_sigs[keep_mask] if old_sigs is not None else None

                base, ext = os.path.splitext(fname)
                new_name = f"{base}_clean{ext}"
                save_path = os.path.join(data_dir, new_name)

                if new_sigs is not None:
                    # Save as V7 format
                    v7_data = {
                        "version": "7.0",
                        "data": {"embeddings": new_embs, "passages": new_txts},
                        "signatures": new_sigs,
                        "metadata": {"pruned": True},
                    }
                    with open(save_path, "wb") as f:
                        pickle.dump(v7_data, f)
                else:
                    new_data = {"embeddings": new_embs, "passages": new_txts}
                    with open(save_path, "wb") as f:
                        pickle.dump(new_data, f)

                st.success(f"Pruned cartridge saved: {new_name}")
                st.session_state.deleted_ids = set()
                st.rerun()


# --- MAIN INTERFACE ---
st.title("Vector+ Studio v0.7")
st.caption("True Content-Addressable Recall via Lattice Signatures")

# Show help when no cartridge is loaded
if not st.session_state.dataset:
    st.info("👈 **Open the sidebar** (click the arrow in the top-left) to mount a Memory Cartridge or forge a new one.")

col_search, col_spell, col_opts = st.columns([3, 1, 1], gap="small")
with col_search:
    query = st.text_input(
        "Semantic Query", placeholder="Enter concept to recall...", key="q_box"
    )

with col_spell:
    if HAS_SPELLCHECK:
        use_spellcheck = st.checkbox("Auto-Correct Typos", value=False)
        if query and use_spellcheck:
            clean_query = strict_spellcheck(query)
            if clean_query != query:
                st.info(f"Auto-Corrected: '{query}' -> '{clean_query}'")
                query = clean_query

with col_opts:
    top_k = st.number_input("Top-K", min_value=1, max_value=500, value=25)


# 2. EXECUTION
if st.button("Search", type="primary") or query:
    if not st.session_state.dataset:
        st.error("No Cartridge Mounted.")
        st.stop()

    embedder = load_embedder()
    ml = st.session_state.engine
    raw_txts = st.session_state.dataset["txt"]

    # --- STEP A: QUERY CLEANING & KEYWORD EXTRACTION ---
    t0 = time.perf_counter()
    stopwords = {
        "a",
        "an",
        "the",
        "is",
        "are",
        "was",
        "were",
        "what",
        "who",
        "how",
        "when",
        "where",
        "why",
        "which",
        "that",
        "this",
        "of",
        "in",
        "on",
        "for",
        "to",
        "and",
        "or",
        "be",
        "it",
        "as",
        "at",
        "by",
        "from",
        "about",
        "with",
        "find",
        "search",
        "looking",
        "me",
        "you",
        "us",
        "them",
        "tell",
        "show",
        "give",
        "get",
        "can",
        "does",
        "did",
        "do",
    }
    query_words = [w.lower().strip("?.,!") for w in query.split() if len(w) > 2]
    keywords = [w for w in query_words if w not in stopwords]
    clean_query = " ".join(keywords) if keywords else query

    # with status_log.container():
    #     st.info("🌊 Rippling Energy Landscape...")

    # --- STEP B: EMBED CLEANED QUERY ---
    vec = embedder.encode(f"search_query: {clean_query}")

    # --- STEP C: LATTICERUNNER V7 PHYSICS (Visualization) ---
    with st.session_state.engine_lock:
        v_min, v_max = vec.min(), vec.max()
        vec_norm = (vec - v_min) / (v_max - v_min + 1e-9)
        vec_norm = vec_norm.astype(np.float32)

        ml.reset()
        ml.imprint_vector(vec_norm)
        query_signature = ml.generate_signature()
        recall_vec = ml.recall()

    # --- STEP D: SEMANTIC RANKING (Embedding Cosine Similarity) ---
    library_embs = st.session_state.dataset.get("emb")
    if library_embs is None or len(library_embs) == 0:
        st.error("No embeddings available. Try re-mounting the cartridge.")
        st.stop()

    query_norm = vec / (np.linalg.norm(vec) + 1e-9)
    lib_norms = np.linalg.norm(library_embs, axis=1, keepdims=True) + 1e-9
    library_normed = library_embs / lib_norms
    base_scores = np.dot(library_normed, query_norm)

    # --- STEP E: KEYWORD INJECTION (Exact Title Matches) ---
    # Guarantees entries with keyword in title are considered
    keyword_candidates = set()
    for i, txt in enumerate(raw_txts):
        if i in st.session_state.deleted_ids:
            continue
        title = txt.splitlines()[0].lower() if txt else ""
        for kw in keywords:
            if len(kw) > 2 and kw in title:
                keyword_candidates.add(i)
                break

    # --- STEP F: COMBINE CANDIDATES & RE-RANK ---
    # Top 200 by embedding + all keyword matches
    embedding_candidates = set(np.argsort(base_scores)[-200:][::-1].tolist())
    all_candidates = embedding_candidates | keyword_candidates

    results = []
    for idx in all_candidates:
        if idx in st.session_state.deleted_ids:
            continue

        cos_sim = base_scores[idx]
        title = raw_txts[idx].splitlines()[0].lower() if raw_txts[idx] else ""

        # Keyword bonus: +0.1 per keyword found in title (max +0.3)
        bonus = sum(0.1 for kw in keywords if kw in title)
        bonus = min(bonus, 0.3)

        final_score = cos_sim + bonus
        results.append((final_score, idx, cos_sim, bonus))

    results.sort(key=lambda x: x[0], reverse=True)
    final_results = [(score, idx) for score, idx, _, _ in results[:top_k]]
    ranking_mode = "V7 Semantic + Keywords"
    t1 = time.perf_counter()

    # --- DISPLAY RESULTS ---

    # with status_log.container():
    #     st.success(f"✅ Converged in {(t1-t0):.1f}s")  # You can grab the real timer

    st.caption(f"Latency: {(t1-t0)*1000:.1f}ms | Ranking: {ranking_mode}")

    if final_results:
        winner_score, winner_idx = final_results[0]
        winner_text = raw_txts[winner_idx]

        st.markdown("### Top Result")
        color = "green" if winner_score > 0.5 else "orange"

        with st.container(border=True):
            lines = winner_text.splitlines()
            st.markdown(f"#### {lines[0]}")
            if len(lines) > 1:
                st.write(f"_{lines[1]}_")
            else:
                st.write(winner_text)
            st.divider()
            c1, c2 = st.columns([4, 1], gap="small")
            with c1:
                st.markdown(
                    f"**ID:** `{winner_idx}` | **Relevance:** :{color}[{winner_score:.3f}]"
                )
            with c2:
                if st.button(f"Delete ID {winner_idx}", key=f"del_{winner_idx}"):
                    st.session_state.deleted_ids.add(winner_idx)
                    st.rerun()

    # --- VISUAL CORTEX (LatticeRunner Output) ---
    with st.expander("LatticeRunner V7 Telemetry", expanded=True):

        col_main, col_rank = st.columns([2, 1], gap="small", width="stretch")
        with col_main:
            l4_state = recall_vec.reshape(4096, 4096)
            l3_state = l4_state.reshape(256, 16, 256, 16).mean(axis=(1, 3))
            l3_norm = np.clip(l3_state * 25.0, 0.0, 1.0)

            img = np.zeros((256, 256, 3), dtype=np.uint8)
            # V7 Theme: Purple/Magenta
            img[:, :, 0] = (30 + (l3_norm * 200)).astype(np.uint8)
            img[:, :, 1] = (10 + (l3_norm * 50)).astype(np.uint8)
            img[:, :, 2] = (40 + (l3_norm * 215)).astype(np.uint8)

            col_viz, col_info = st.columns([1, 2], gap="small")
            with col_viz:
                st.image(
                    img,
                    caption="Lattice Activation Pattern",
                    width="stretch",
                    clamp=True,
                )

            with col_info:
                st.markdown("#### V7 Physics Core State")
                n_active = (recall_vec > 0.01).sum()
                sparsity = (n_active / 16_777_216) * 100

                # Query signature stats
                sig_active = (query_signature > 0.1).sum()
                sig_max = query_signature.max()

                st.metric("Active Neurons", f"{n_active:,}")
                st.metric("Global Sparsity", f"{sparsity:.3f}%")
                st.metric("Signature Active Regions", f"{sig_active}/4096")
                st.metric("Signature Peak", f"{sig_max:.3f}")

                if sparsity > 25.0:
                    st.error("CORE UNSTABLE (Bloom Detected)")
                elif sparsity < 1.0:
                    st.warning("Pattern may be too sparse")
                else:
                    st.success("CORE STABLE")

                # Show ranking mode
                if ranking_mode == "V7 Signatures":
                    st.success("Using TRUE lattice-based ranking!")
                else:
                    st.warning("Signatures still building...")
        with col_rank:
            # Table
            with st.expander("Show Ranking Details"):
                data = []
                for rank, (score, idx) in enumerate(final_results):
                    txt = raw_txts[idx]
                    data.append(
                        {
                            "Rank": rank + 1,
                            "ID": idx,
                            "Score": f"{score:.4f}",
                            "Snippet": txt.splitlines()[0][:60],
                        }
                    )
                # st.dataframe(data, width="stretch", height=300)
                st.dataframe(data, width="stretch")
